# extractor.py
import io
import pdfplumber
from docx import Document
from fastapi import UploadFile, HTTPException

async def extract_text(file: UploadFile) -> str:
    filename = file.filename.lower()
    content = await file.read()

    if filename.endswith(".pdf"):
        return _extract_pdf(content)
    elif filename.endswith(".docx"):
        return _extract_docx(content)
    else:
        raise HTTPException(
            status_code=400,
            detail="Only PDF and DOCX resumes are supported."
        )


def _extract_pdf(content: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            # layout-aware extraction — handles multi-column resumes
            # (common template style) better than raw pypdf
            page_text = page.extract_text(layout=True)
            if page_text:
                text_parts.append(page_text)

    text = "\n".join(text_parts).strip()

    if not text:
        # No text layer at all → likely a scanned/image-based resume
        raise HTTPException(
            status_code=422,
            detail="Couldn't extract text — this looks like a scanned/image PDF, not a text-based one."
        )

    return text


def _extract_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = []

    # paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # tables (some resume templates put content in tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()

    if not text:
        raise HTTPException(status_code=422, detail="This DOCX appears to be empty.")

    return text